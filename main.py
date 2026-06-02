from collections import Counter, defaultdict
from pathlib import Path
import re
import sys

from docx import Document  # type: ignore
from openpyxl import Workbook  # type: ignore
from openpyxl.styles import Border, Side

BASE_DIR = Path(__file__).resolve().parent
INPUT_DIR = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else BASE_DIR / "Descontar"
OUTPUT_FILE = Path(sys.argv[2]).resolve() if len(sys.argv) > 2 else INPUT_DIR / "resumen_productos.xlsx"

PRODUCT_PATTERN = re.compile(r"^\s*(\d+)\s+(.+?)\s*$", re.IGNORECASE)
NOISE_PATTERN = re.compile(
    r"\b(?:DIRECCION|DIRECCIÓN|TELÉFONO|TELEFONO|EMAIL|E-MAIL|DNI|PEDIDO|CANJE|MENSAJERÍA|MENSAJERIA|ANDREANI|MERCADO ENVIOS|ENVIOS FLEX|CASA|DEPARTAMENTO|AV\.|CALLE|BARRIO|RUTA|SARGENTO|PA\b|LOMAS|BUENOS AIRES|RÍO|RIO|DE JULIO|DE MAYO|DE JUNIO)\b",
    re.IGNORECASE,
)


def normalize_product(name: str) -> str:
    """Normaliza el texto del producto para agrupar repeticiones."""
    text = re.sub(r"\s+", " ", name).strip()
    return text.title() if text else ""


def is_candidate_product_line(text: str) -> bool:
    raw = text.strip()
    if not raw:
        return False

    if NOISE_PATTERN.search(raw):
        return False

    match = PRODUCT_PATTERN.match(raw)
    if not match:
        return False

    quantity = int(match.group(1))
    product = normalize_product(match.group(2))
    if quantity <= 0 or len(product) < 4:
        return False

    return True


def iter_docx_text(file_path: Path):
    """Recoge texto línea por línea para capturar todos los productos dentro de cada párrafo."""
    doc = Document(str(file_path))

    for paragraph in doc.paragraphs:
        for line in paragraph.text.splitlines():
            raw = line.strip()
            if raw:
                yield raw

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for line in paragraph.text.splitlines():
                        raw = line.strip()
                        if raw:
                            yield raw


def extract_products_from_docx(file_path: Path):
    """Lee un DOCX y devuelve una lista de (cantidad, producto)."""
    products = []

    for raw in iter_docx_text(file_path):
        if not is_candidate_product_line(raw):
            continue

        match = PRODUCT_PATTERN.match(raw)
        if not match:
            continue

        assert match is not None
        quantity = int(match.group(1))
        product = normalize_product(match.group(2))

        if quantity <= 0 or not product:
            continue

        products.append((quantity, product))

    return products


def build_summary():
    if not INPUT_DIR.exists() or not INPUT_DIR.is_dir():
        raise FileNotFoundError(f"No se encontró la carpeta de entrada: {INPUT_DIR}")

    totals = Counter()
    files_used = defaultdict(int)

    for file_path in sorted(INPUT_DIR.glob("*.docx")):
        for quantity, product in extract_products_from_docx(file_path):
            totals[product] += quantity
            files_used[product] += 1

    return totals, files_used


def save_excel(totals: Counter, files_used: defaultdict):
    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Resumen"
    ws.append(["Producto", "Cantidad total"])

    for product, quantity in sorted(totals.items(), key=lambda item: (-item[1], item[0])):
        ws.append([product, quantity])

    thin = Side(border_style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=2):
        for cell in row:
            cell.border = border

    ws.column_dimensions["A"].width = 80
    ws.column_dimensions["B"].width = 18
    ws.freeze_panes = "A2"

    wb.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    totals, files_used = build_summary()
    output_file = save_excel(totals, files_used)
    print(f"Carpeta leída: {INPUT_DIR}")
    print(f"Resumen generado en: {output_file}")
    print("\nProductos totales:")
    for product, quantity in sorted(totals.items(), key=lambda item: (-item[1], item[0])):
        print(f"- {quantity:>3}  {product}")
    input("\nPresiona Enter para salir...")